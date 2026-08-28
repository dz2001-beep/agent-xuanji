#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""生成《AI Agent（智能体）研究报告》Word 文档"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

# 中文字体设置函数
def set_cn_font(run, name="微软雅黑", size=None, bold=None, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color

doc = Document()

# ============ 全局样式 ============
style = doc.styles["Normal"]
style.font.name = "微软雅黑"
style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
style.font.size = Pt(11)

# 页边距
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

# ============ 封面标题 ============
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_before = Pt(120)
run = title.add_run("\nAI Agent（智能体）技术研究报告")
set_cn_font(run, size=26, bold=True, color=RGBColor(0x1F, 0x4E, 0x79))

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run("从概念原理到产业应用的全面解析")
set_cn_font(run, size=15, color=RGBColor(0x59, 0x59, 0x59))

# 分隔线
sep = doc.add_paragraph()
sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sep.add_run("—" * 30)
set_cn_font(run, size=12, color=RGBColor(0x1F, 0x4E, 0x79))

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta.add_run("编制日期：2026年\n保密等级：内部资料")
set_cn_font(run, size=12, color=RGBColor(0x80, 0x80, 0x80))

doc.add_page_break()

# ============ 辅助函数 ============
def h1(text):
    p = doc.add_heading(level=1)
    run = p.add_run(text)
    set_cn_font(run, size=18, bold=True, color=RGBColor(0x1F, 0x4E, 0x79))
    return p

def h2(text):
    p = doc.add_heading(level=2)
    run = p.add_run(text)
    set_cn_font(run, size=14, bold=True, color=RGBColor(0x2E, 0x74, 0xB5))
    return p

def para(text, bold=False, size=11, italic=False, color=None, indent=0):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.4
    run = p.add_run(text)
    set_cn_font(run, size=size, bold=bold, color=color)
    return p

def bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    set_cn_font(run, size=11)
    return p

def numbered(text):
    p = doc.add_paragraph(style="List Number")
    run = p.add_run(text)
    set_cn_font(run, size=11)
    return p

# ============ 目录页 ============
h1("目录")
toc_items = [
    "1. 引言",
    "2. 什么是 AI Agent",
    "   2.1 定义",
    "   2.2 核心特征",
    "3. AI Agent 的系统架构",
    "4. AI Agent 的关键技术与能力",
    "5. 主流框架与产品",
    "6. 应用场景",
    "7. 面临的挑战与风险",
    "8. 未来发展趋势",
    "9. 总结与展望",
]
for item in toc_items:
    p = doc.add_paragraph()
    run = p.add_run(item)
    set_cn_font(run, size=12)

doc.add_page_break()

# ============ 1. 引言 ============
h1("1. 引言")
para("近年来，随着大语言模型（LLM）技术的爆发式发展，人工智能领域正经历一次范式转型——从单纯的“回答问题”走向“采取行动”。以 ChatGPT、Claude 等为代表的对话式模型，其能力边界被不断拓展，而“智能体（Agent）”正成为将模型能力转化为真实世界行动力的关键载体。", indent=True)
para("本报告旨在系统梳理 AI Agent 的概念定义、核心技术、系统架构、主流框架、典型应用与未来趋势，帮助读者建立对智能体技术的全面认知。", indent=True)

# ============ 2. 什么是 AI Agent ============
h1("2. 什么是 AI Agent")

h2("2.1 定义")
para("AI Agent（人工智能智能体）是一种能够通过感知环境、自主推理、规划决策并执行行动来达成特定目标的智能系统。它不同于传统的一次性问答模型，而是具备“目标导向 + 持续迭代 + 工具调用 + 环境反馈”能力的主动性实体。", indent=True)
para("业界普遍认可的一个形象比喻是：模型好比“大脑”，而 Agent 则是在大脑支配下能够“动手做事”的完整个体。", indent=True)

h2("2.2 核心特征")
bullet("自主性：在无人干预或少量干预下，自主规划并执行任务。")
bullet("目标导向：围绕用户给定的目标，拆解子任务并逐步推进。")
bullet("工具使用：能够调用 API、搜索引擎、代码执行器、数据库等外部工具。")
bullet("记忆能力：具备短期上下文记忆与长期知识存储，支撑连贯性推理。")
bullet("环境交互：通过与环境（软件、网页、物理设备）的反馈不断调整策略。")

# ============ 3. 系统架构 ============
h1("3. AI Agent 的系统架构")
para("一个典型的 AI Agent 系统通常由以下核心模块组成：", indent=True)
architecture = [
    ("规划模块（Planning）", "负责将复杂目标分解为可执行的子任务，并设计执行序列。"),
    ("记忆模块（Memory）", "包含短期工作记忆与长期记忆（向量数据库、知识图谱），用于存储上下文和历史经验。"),
    ("工具模块（Tools）", "封装外部能力，如代码解释器、网页抓取、API 网关、数据库查询、第三方服务调用。"),
    ("反思模块（Reflection）", "对执行结果进行自我评估与修正，支持“试错—改进”的自适应循环。"),
    ("执行模块（Action）", "与环境进行实际交互，产生可观察的行为或输出。"),
]
for name, desc in architecture:
    p = doc.add_paragraph()
    run = p.add_run(f"★ {name}：")
    set_cn_font(run, size=11, bold=True)
    run = p.add_run(desc)
    set_cn_font(run, size=11)

# ============ 4. 关键技术与能力 ============
h1("4. AI Agent 的关键技术与能力")
table_head = ["技术方向", "说明", "代表性方法"]
table_data = [
    ["思考链推理", "引导模型逐步推理，提升复杂问题的解决能力", "Chain-of-Thought (CoT)、Tree of Thoughts"],
    ["工具学习", "让模型学会何时、如何调用外部工具", "Function Calling、ToolFormer"],
    ["任务规划", "将高层次目标拆解为有序步骤", "ReAct、Plan-and-Execute"],
    ["记忆机制", "持久化上下文与经验，支持长程任务", "向量检索（RAG）、内存缓冲"],
    ["多智能体协作", "多个 Agent 分工协作，解决复杂系统性问题", "AutoGen、Multi-Agent Debate"],
    ["自我反思", "基于反馈进行错误修正与策略优化", "Reflexion、Self-Refine"],
    ["强化学习", "通过环境奖励信号优化决策策略", "RLHF、PPO"],
]
table = doc.add_table(rows=1, cols=3)
table.style = "Light Grid Accent 1"
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
for i, t in enumerate(table_head):
    hdr[i].text = ""
    run = hdr[i].paragraphs[0].add_run(t)
    set_cn_font(run, size=11, bold=True)
for row in table_data:
    cells = table.add_row().cells
    for i, val in enumerate(row):
        cells[i].text = ""
        run = cells[i].paragraphs[0].add_run(val)
        set_cn_font(run, size=10)

# ============ 5. 主流框架与产品 ============
h1("5. 主流框架与产品")
bullet("AutoGen（微软）：支持多智能体对话与协作的通用框架。")
bullet("LangChain / LangGraph：提供 Agent 编排、工具调用与工作流管理的生态。")
bullet("LlamaIndex：侧重数据接入与检索增强的 Agent 框架。")
bullet("MetaGPT：面向软件项目的多智能体协作框架。")
bullet("OpenAI Assistants API / Function Calling：面向生产环境的标准智能体接口。")
bullet("Dify、Coze：面向低代码场景的智能体应用构建平台。")

# ============ 6. 应用场景 ============
h1("6. 应用场景")
scene = [
    ("智能客服与个人助理", "自动处理咨询、预约、购物、日程管理等事务。"),
    ("代码开发", "自动写代码、调试、生成测试、代码审查与重构。"),
    ("数据分析与报告", "自动采集数据、清洗、建模分析并生成可视化报告。"),
    ("企业流程自动化（RPA 升级）", "将传统规则自动化升级为理解型、决策型的智能流程。"),
    ("科学研究", "辅助文献调研、实验设计、假设生成与论文写作。"),
    ("金融与医疗", "智能投研、风险评估、辅助诊断与病历分析。"),
    ("游戏与仿真", "打造具有自主决策能力的 NPC 与数字人。"),
]
for name, desc in scene:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(f"{name}：")
    set_cn_font(run, size=11, bold=True)
    run = p.add_run(desc)
    set_cn_font(run, size=11)

# ============ 7. 挑战与风险 ============
h1("7. 面临的挑战与风险")
bullet("可靠性：AI Agent 偶尔会“出错”，在关键任务中缺乏兜底机制。")
bullet("安全性：工具调用可能带来越权操作、数据泄露与提示注入攻击。")
bullet("可解释性：多步骤的自主决策过程难以完全透明可追溯。")
bullet("成本与延迟：复杂任务依赖多次模型调用，计算开销与响应耗时较高。")
bullet("数据与隐私：涉及大量敏感数据，合规与隐私保护要求严格。")
bullet("评估体系：缺乏统一、可量化的评测基准来度量 Agent 的真实能力。")

# ============ 8. 未来趋势 ============
h1("8. 未来发展趋势")
numbered("从“单智能体”走向“多智能体系统”，形成专业化分工与协作生态。")
numbered("从“工具调用”走向“全栈自主”，覆盖感知、决策、执行的完整闭环。")
numbered("从“云端托管”走向“端侧智能”，本地设备上的轻量 Agent 将普及。")
numbered("记忆与个性化深度融合，Agent 将成为具备长期关系的“数字伙伴”。")
numbered("安全、对齐与合规技术将成为 Agent 大规模落地的前置条件。")
numbered("Agent 将逐渐成为企业数字化转型的“超级入口”与“新生产力”。")

# ============ 9. 总结与展望 ============
h1("9. 总结与展望")
para("AI Agent 代表着人工智能从“被动问答”迈向“主动行动”的关键一步。它把大语言模型的推理能力与对外部世界的操作能力进行有机结合，正在重塑软件开发、企业运营、个人生活等多个领域的工作方式。", indent=True)
para("尽管在可靠性、安全性与成本等方面仍面临诸多挑战，但随着多智能体协作、记忆增强、工具生态的不断成熟，Agent 技术将迎来更广泛、更深度的落地应用。可以预见，未来几年智能体将如同今天的应用软件一样，成为数字世界中无处不在的基础设施。", indent=True)

# 落款
doc.add_paragraph()
end = doc.add_paragraph()
end.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = end.add_run("报告完\n（本报告基于公开资料整理，仅供内部参考）")
set_cn_font(run, size=10, color=RGBColor(0x80, 0x80, 0x80))

# 保存
out = "AI_Agent研究报告.docx"
doc.save(out)
print(f"已生成：{out}")
